from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTImage

def extract_text_and_images(pdf_path):
    text = ""
    images = []
    for page_layout in extract_pages(pdf_path):
        for element in page_layout:
            if isinstance(element, LTTextContainer):
                text += element.get_text()
            elif isinstance(element, LTImage):
                # 如果图片是直接嵌入在 PDF 中，你可以使用 element.stream 获取图片数据
                # 如果图片是通过外部链接引用的，你可以使用 element.uri 获取图片链接
                images.append(element.stream)
    return text, images

def split_text_by_empty_lines(text):
    sections = text.split("\n\n")
    return sections

from docx import Document
def save_sections_as_word(sections, filename):
    doc = Document()
    for i, section in enumerate(sections):
        doc.add_heading(f"Section {i+1}", level=1)
        doc.add_paragraph(section)
    doc.save(filename)

# 使用方法
pdf_path = "1.pdf"
text, images = extract_text_and_images(pdf_path)
print(text, images)

sections = split_text_by_empty_lines(text)
save_sections_as_word(sections, "../outputs/pdf_meta.docx")
