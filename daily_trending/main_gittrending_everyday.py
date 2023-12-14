import requests
from bs4 import BeautifulSoup
import json
import re
import os
from datetime import datetime
import pytz

today = datetime.now(pytz.timezone('US/Eastern')).strftime('%Y-%m-%d')

def scrape_github_trending():
    headers = {'User-Agent': 'Mozilla/5.0'}
    url = "https://github.com/trending"
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.content, 'html.parser')

    output = []
    for article in soup.find_all('article', class_="Box-row"):
        repo_name = article.find("h2", {"class": "h3 lh-condensed"}).get_text(strip=True)
        repo_owner, repo_name = repo_name.split('/')
        repo_description = article.find("p", {"class": "col-9 color-fg-muted my-1 pr-4"}).get_text(strip=True)
        
        main_language_element = article.find(itemprop="programmingLanguage")
        main_language = main_language_element.get_text(strip=True) if main_language_element else None

        stars_text = article.find(href=re.compile("/stargazers")).get_text(strip=True)
        if "k" in stars_text:
            stars = float(stars_text[:-1]) * 1000
        else:
            stars = str(stars_text)

        forks_text = article.find(href=re.compile("/forks")).get_text(strip=True)
        if "k" in forks_text:
            forks = float(forks_text[:-1]) * 1000
        else:
            forks = str(forks_text)

        built_by = [img.get('alt') for img in article.find_all('img', {'class': 'avatar mb-1 avatar-user'})]
        repo_link = article.h2.find('a').get('href')
        if repo_link is not None:
            repo_link_re = "https://github.com" + repo_link
        else:
            repo_link_re = None
        
        repo_info = {
            "repo_name": repo_name,
            "repo_owner": repo_owner,
            "repo_description": repo_description,
            "main_language": main_language,
            "stars": stars,
            "forks": forks,
            "built_by": built_by,
            "repo_link": repo_link_re,
            "date": today
        }
        output.append(repo_info)

    # Load existing data from json file (if it exists)
    json_file = './data/data_trending/github_trending.json'
    if os.path.exists(json_file):
        with open(json_file, 'r') as f:
            existing_data = json.load(f)
        
        # Merge existing data with new data and remove duplicates by repo name (keeping the earlier date)
        merged_data = existing_data + output
        merged_data.sort(key=lambda x: x['date'], reverse=True)
        unique_data = []
        for item in merged_data:
            if item['repo_name'] not in [d['repo_name'] for d in unique_data]:
                unique_data.append(item)
        
        output = unique_data

    # Save data to json file
    with open(json_file, 'w') as f:
        json.dump(output, f, indent=4)

from operator import itemgetter

from docx import Document

def json_to_word(json_file):
    with open(json_file, 'r') as jf:
        data = json.load(jf)
    today_repos = [repo for repo in data if repo['date'] == today]
    today_repos.sort(key=itemgetter('stars'), reverse=True)
    daty = today.replace("-", "")
    doc = Document()
    doc.add_heading("Github Trending 今日热门项目", 0)
    for i, repo in enumerate(today_repos):
        doc.add_heading(f"{i+1}. {repo['repo_name']} by {repo['repo_owner']}", level=1)
        doc.add_paragraph(f"Description: {repo['repo_description']}")
        doc.add_paragraph(f"Main Language: {repo['main_language']}")
        doc.add_paragraph(f"Stars: {repo['stars']}")
        doc.add_paragraph(f"Forks: {repo['forks']}")
        doc.add_paragraph(f"Link: {repo['repo_link']}")
        doc.add_paragraph(f"Built by: {', '.join(repo['built_by'])}")
        doc.add_paragraph("---")
    doc.save(f'./outputs/github_trending_{daty}.docx')
def json_to_markdown(json_file):
    with open(json_file, 'r') as jf:
        data = json.load(jf)
    today_repos = [repo for repo in data if repo['date'] == today]
    today_repos.sort(key=itemgetter('stars'), reverse=True)
    daty = today.replace("-", "")
    md_file = f'./outputs/github_trending_{daty}.md'
    with open(md_file, 'w') as mf:
        mf.write("Github Trending 今日热门项目\n")
        for i, repo in enumerate(today_repos):
            mf.write(f"### {i+1}. [{repo['repo_name']}]({repo['repo_link']}) by {repo['repo_owner']} 📚\n")
            mf.write(f"📖 Description:{repo['repo_description']}\n")
            mf.write(f"🔭 Main Language:{repo['main_language']}\n")
            mf.write(f"⭐ Stars:{repo['stars']}\n")
            mf.write(f"🍴 Forks:{repo['forks']}\n")
            mf.write(f"🔨 Built by:{', '.join(repo['built_by'])}\n")
            mf.write("---\n")

if not os.path.exists('./data/data_trending'):
    os.makedirs('./data/data_trending')
if not os.path.exists('./outputs'):
    os.makedirs('./outputs')
scrape_github_trending()
json_to_word('./data/data_trending/github_trending.json')
#json_to_markdown('./data/data_trending/github_trending.json')
