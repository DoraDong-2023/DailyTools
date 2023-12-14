import json
import os
import arxiv
from datetime import datetime
import requests
from utils.eng2ch import eng2zh
import pytz


def arxiv_main(today):
    physics_topics = ['physics.atom-ph', 'physics.class-ph', 'physics.comp-ph', 'physics.data-an', 'physics.flu-dyn', 'physics.gen-ph', 'physics.geo-ph', 'physics.hist-ph', 'physics.ins-det']
    math_topics = ['math.AG', 'math.AT', 'math.AP', 'math.CT', 'math.CA', 'math.CO', 'math.AC', 'math.CV', 'math.DG', 'math.DS', 'math.FA', 'math.GM', 'math.GN', 'math.GT', 'math.GR', 'math.HO', 'math.IT', 'math.KT', 'math.LO', 'math.MP', 'math.MG', 'math.NT', 'math.NA', 'math.OA', 'math.OC', 'math.PR', 'math.QA', 'math.RT', 'math.RA', 'math.SP', 'math.ST', 'math.SG']
    bio_topics = ['q-bio.BM', 'q-bio.CB', 'q-bio.GN', 'q-bio.MN', 'q-bio.NC', 'q-bio.OT', 'q-bio.PE', 'q-bio.QM', 'q-bio.SC', 'q-bio.TO']
    cs_topics = ['cs.AR', 'cs.AI', 'cs.CL', 'cs.CC', 'cs.CE', 'cs.CG', 'cs.GT', 'cs.CV', 'cs.CY', 'cs.CR', 'cs.DS', 'cs.DB', 'cs.DL', 'cs.DM', 'cs.DC', 'cs.GL', 'cs.GR', 'cs.HC', 'cs.IR', 'cs.IT', 'cs.LG', 'cs.LO', 'cs.MS', 'cs.MA', 'cs.MM', 'cs.NI', 'cs.NE', 'cs.NA', 'cs.OS', 'cs.OH', 'cs.PF', 'cs.PL', 'cs.RO', 'cs.SE', 'cs.SD', 'cs.SC', 'cs.SY']
    fin_topics = ['q-fin.CP', 'q-fin.GN', 'q-fin.MF', 'q-fin.PM', 'q-fin.PR', 'q-fin.RM', 'q-fin.ST', 'q-fin.TR']
    stat_topics = ['stat.AP', 'stat.CO', 'stat.ML', 'stat.ME', 'stat.OT', 'stat.TH']
    #topics = bio_topics + cs_topics + physics_topics + math_topics + fin_topics + stat_topics

    topics = ['cs.LG', 'cs.AI', 'cs.CL', 'cs.NE', 'cs.RO'] + \
            bio_topics+ \
            stat_topics + \
            fin_topics
            
    base_dir = './data/arxiv'

    download_flag = False

    topic_results = {}  # Save the results of each prefix

    for topic in topics:
        prefix = topic.split(".")[0]
        dirpath = os.path.join(base_dir, prefix)
        
        search = arxiv.Search(
            query=topic,
            max_results=100,
            sort_by=arxiv.SortCriterion.SubmittedDate,
            sort_order=arxiv.SortOrder.Descending,
        )

        results_today = []
        for result in search.results():
            if str(result.published.date()).strip() == today:
                paper_dict = {
                    'title': result.title,
                    'authors': [str(author) for author in result.authors],
                    'summary': result.summary,
                    'published': str(result.published),
                    'pdf_url': result.pdf_url,
                    'downloaded_time': today.strip()
                }
                results_today.append(paper_dict)
            else:
                break

            if download_flag:
                if 'pdf' in result.pdf_url:
                    filename = os.path.join(dirpath, os.path.basename(result.pdf_url)+'.pdf')
                    response = requests.get(result.pdf_url, stream=True)
                    with open(filename, 'wb') as f:
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                f.write(chunk)
                else:
                    print(f'Unknown file {result.pdf_url}. Skipping download.')

        if prefix not in topic_results:
            topic_results[prefix] = []
        topic_results[prefix].extend(results_today)

    # Save each prefix's papers to a separate json file
    for prefix, results in topic_results.items():
        file_path = os.path.join(base_dir, f"{prefix}_papers.json")
        with open(file_path, 'w') as f:
            json.dump(results, f)

import json
from docx import Document

def generate_markdown(data):
    md_content = ""
    for paper in data:
        md_content += f"# {paper['title']}\n"
        md_content += f"**Authors:** {', '.join(paper['authors'])}\n"
        md_content += f"**Summary:**\n{paper['summary']}\n"
        md_content += f"**Published:** {paper['published']}\n"
        md_content += f"**PDF URL:** [Link]({paper['pdf_url']})\n"
        md_content += f"**Downloaded Time:** {paper['downloaded_time']}\n\n"
    with open('./outputs/output.md', 'w') as f:
        f.write(md_content)

def generate_docx(data,title):
    doc = Document()
    for paper in data:
        doc.add_heading(paper['title'], level=1)
        doc.add_paragraph("Authors: " + ', '.join(paper['authors']))
        doc.add_paragraph("Summary:\n" + paper['summary'])
        doc.add_paragraph("Published: " + paper['published'])
        doc.add_paragraph("PDF URL: " + paper['pdf_url'])
        doc.add_paragraph("Downloaded Time: " + paper['downloaded_time'])
        doc.add_paragraph("\n")
    doc.save(f'./outputs/arxiv_trending_{title}_{today}.docx')

def generate_iterative_by_docx(directory):
    for filename in os.listdir(directory):
        if filename.endswith('.json'):
            with open(os.path.join(directory, filename), 'r') as f:
                file_data = json.load(f)
                generate_docx(file_data,filename.replace('.json',''))
    

if __name__ == '__main__':
    today = datetime.now(pytz.timezone('US/Eastern')).date()
    today = '2023-07-31'
    arxiv_main(today)
    base_dir = './data/arxiv'
    articles_all = generate_iterative_by_docx('./data/arxiv')
    #generate_markdown(articles_all)
    generate_docx(articles_all)