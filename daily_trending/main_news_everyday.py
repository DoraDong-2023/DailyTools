import requests
import json
from datetime import datetime
import pytz
import os
import markdown
import re
from bs4 import BeautifulSoup

today = datetime.now(pytz.timezone('US/Eastern')).strftime('%Y-%m-%d')
#today = '2023-07-20'
today_save = today.replace("-", "")
sources_list = [
    "abc-news","al-jazeera-english","ars-technica","associated-press","australian-financial-review",
    "bbc-news","bloomberg","business-insider","cnn","engadget","fortune","fox-news","google-news",
    "hacker-news","ign","independent","info-money","medical-news-today","msnbc","national-geographic","nbc-news",
    "newsweek","new-scientist","next-big-future","politico","reuters","techcrunch","techradar","the-guardian-uk",
    "the-new-york-times","the-verge","the-wall-street-journal","the-washington-post","time","wired"
]

query_terms = ['artificial intelligence', 'machine learning', 'quantum computing', 'blockchain', 'cybersecurity', 'augmented reality', 'virtual reality', 'biotechnology', 'fintech', 'startups funding']

def get_news_articles(api_key, query_terms):
    url = "https://newsapi.org/v2/everything"
    news_dict = {}
    if not os.path.exists('./data/data_trending'):
        os.makedirs('./data/data_trending')

    articles_all = []
    for source in sources_list:
        news_dict[source] = {}
        for term in query_terms:
            parameters = {
                'q': term,
                'sources': source,
                'sortBy': 'popularity',
                'language': 'en',
                'from': today,
                'to': datetime.now(pytz.timezone('US/Eastern')).strftime('%Y-%m-%d'),
                'apiKey': api_key
            }
            try:
                response = requests.get(url, params=parameters)
                data = response.json()
                for article in data['articles']:
                    article['date'] = today
                    #article['content'] = get_full_content(article['url'])
                articles_all = articles_all + data['articles']
                # 筛选今天的数据
                news_dict[source][term] = articles_all
            except Exception as e:
                print(f"Error occurred with source {source} and term {term}: {e}")
    with open(f'./data/data_trending/news.json', 'w') as f:
        json.dump(articles_all, f, indent=4)

def generate_md_file():
    if not os.path.exists('./outputs'):
        os.makedirs('./outputs')
    with open('./data/data_trending/news.json', 'r') as f:
        data = json.load(f)
    with open(f"./outputs/article_{today_save}.md", "w", encoding="utf-8") as f:
        for source in data.keys():
            for term in data[source].keys():
                articles = data[source][term]
                for i, article in enumerate(articles):
                    f.write(f"## {article['title']}\n\n")
                    f.write(f"作者 😎：{article['author']}\n\n")
                    f.write(f"发布日期 📆：{article['publishedAt'].split('T')[0]}\n\n")
                    f.write(f"### 摘要 💬\n\n")
                    f.write(f"{article['description']}\n\n")
                    f.write(f"\n\n---\n\n本文来源 🔗：[{article['source']['name']}]({article['url']})\n\n")

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_word_file():
    if not os.path.exists('./outputs'):
        os.makedirs('./outputs')
    with open('./data/data_trending/news.json', 'r') as f:
        data = json.load(f)
    doc = Document()
    for source in data.keys():
        for term in data[source].keys():
            articles = data[source][term]
            for i, article in enumerate(articles):
                title = doc.add_heading(level=1)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                title_run = title.add_run(article['title'])
                title_run.bold = True
                title_run.italic = True
                doc.add_paragraph(f"Author: {article['author']}")
                doc.add_paragraph(f"Published Date: {article['publishedAt'].split('T')[0]}")
                doc.add_paragraph(f"Summary: {article['description']}")
                doc.add_paragraph(f"Source: {article['source']['name']}")
                doc.add_paragraph(f"Link: {article['url']}")
                # add a line break after each article
                doc.add_paragraph("\n\n")
    doc.save(f"./outputs/article_{today_save}.docx")

def get_full_content(url):
    try:
        response = requests.get(url)
        page_content = BeautifulSoup(response.content, "html.parser")
        print(page_content)
        article_body = page_content.find("div", attrs={"class": "article-body"})
        if article_body:
            return article_body.get_text(strip=True)
        else:
            return None
    except Exception as e:
        print(f"Error occurred while getting full content: {e}")
        return None

if __name__ == '__main__':
    api_key = 'e46427ffc14442ae8f767b30e7b06b03'  # 你的 NewsAPI 密钥
    get_news_articles(api_key, query_terms)
    #generate_md_file()
    generate_word_file()
